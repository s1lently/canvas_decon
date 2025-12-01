"""Homework automation - fetch, LLM, submit"""
import os, sys, re, json, time, mimetypes
from pathlib import Path
from urllib.parse import urlparse
from html import unescape

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
import config
from func import utilPromptFiles
from core.log import log

TARGET_URL = "https://psu.instructure.com/courses/2418560/assignments/17474475"


def _session():
    """Create authenticated session"""
    import requests
    cookies = {c['name']: c['value'] for c in json.load(open(config.COOKIES_FILE))}
    s = requests.Session()
    s.cookies.update(cookies)
    s.headers.update({'User-Agent': 'Mozilla/5.0', 'Accept': 'application/json+canvas-string-ids'})
    return s


def _parse_url(url):
    """Extract course_id, assignment_id from URL"""
    parts = urlparse(url).path.split('/')
    return parts[parts.index('courses')+1], parts[parts.index('assignments')+1]


def get_homework_details(url=None):
    """Fetch assignment + submission from Canvas API"""
    url = url or TARGET_URL
    course_id, assign_id = _parse_url(url)
    s = _session()
    base = f"{config.CANVAS_BASE_URL}/api/v1/courses/{course_id}/assignments/{assign_id}"

    assignment = s.get(base, timeout=20).json()
    submission = s.get(f"{base}/submissions/self", timeout=20).json()

    log.info(f"Assignment: {assignment['name']}")
    return {
        'assignment': {k: assignment.get(k) for k in ['id','name','description','due_at','points_possible','submission_types','html_url']},
        'submission': {k: submission.get(k) for k in ['id','submitted_at','workflow_state','score','grade','late','missing']}
    }


def ask_llm_with_pdfs(description, product, model, prompt, ref_files=[]):
    """Call LLM with prompt + description"""
    personal = ""
    try:
        info = json.load(open(config.PERSONAL_INFO_FILE))
        personal = f"\n**Personal Info:**\n- Name: {info['name']}\n- Age: {info['age']}\n- Weight: {info['weight_kg']}kg\n- Height: {info['height_cm']}cm\n- Location: {info['location']}\n"
    except (IOError, json.JSONDecodeError, KeyError):
        pass

    full_prompt = f"{prompt}\n{personal}\n\n**Description:**\n{description}"
    log.info(f"Calling {product} {model}...")
    return utilPromptFiles.call_ai(full_prompt, product, model, ref_files)


def parse_img_requests(text):
    """Extract [gen_img] blocks, return (clean_text, requests)"""
    requests = []
    for m in re.finditer(r'\[gen_img\]\s*\{([^}]+)\}', text, re.DOTALL):
        content = m.group(1)
        name = re.search(r'name:\s*(.+?)(?:\n|$)', content)
        desc = re.search(r'des:\s*(.+?)(?:\n|$)', content, re.DOTALL)
        if name and desc:
            requests.append({'name': name.group(1).strip(), 'description': desc.group(1).strip()})

    clean = re.sub(r'\[gen_img\]\s*\{[^}]+\}', '', text, flags=re.DOTALL).strip()
    return clean, requests


def generate_images(img_requests, output_dir):
    """Generate images using Gemini"""
    if not img_requests:
        return

    from google import genai
    from PIL import Image
    from io import BytesIO

    client = genai.Client(api_key=config.GEMINI_API_KEY)
    log.info(f"Generating {len(img_requests)} image(s)...")

    for req in img_requests:
        try:
            log.info(f"  - {req['name']}")
            resp = client.models.generate_content(model="gemini-2.5-flash-image", contents=[req['description']])
            for part in resp.candidates[0].content.parts:
                if part.inline_data:
                    img_path = os.path.join(output_dir, req['name'])
                    Image.open(BytesIO(part.inline_data.data)).save(img_path)
                    log.info(f"    Generated: {img_path}")
                    break
        except Exception as e:
            log.error(f"    Failed: {e}")


def md_to_docx(output_path, md_path):
    """Convert markdown to docx"""
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    for line in open(md_path, 'r', encoding='utf-8'):
        line = line.rstrip()
        if not line.strip() or line.strip() in ['[yes]','[no]'] or line.startswith('```'):
            continue

        # Headers
        if line.startswith('### '):
            doc.add_heading(line[4:], 3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], 2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], 1)
        # Lists
        elif line.startswith(('*   ', '- ')):
            text = line[4:] if line.startswith('*   ') else line[2:]
            p = doc.add_paragraph(style='List Bullet')
            _add_bold_text(p, text)
        # Indented
        elif line.startswith('    '):
            p = doc.add_paragraph()
            _add_bold_text(p, line.strip())
            p.paragraph_format.left_indent = Inches(0.5)
        # Normal
        else:
            p = doc.add_paragraph()
            _add_bold_text(p, line)

    doc.save(output_path)
    log.info(f"Converted: {output_path}")


def _add_bold_text(para, text):
    """Add text with **bold** support"""
    if '**' not in text:
        para.add_run(text)
        return
    for i, part in enumerate(text.split('**')):
        run = para.add_run(part)
        run.bold = (i % 2 == 1)


def submit_to_canvas(url=None, assignment_folder=None):
    """Upload files and submit to Canvas"""
    import requests as req

    url = url or TARGET_URL
    course_id, assign_id = _parse_url(url)
    output_dir = os.path.join(assignment_folder, 'auto', 'output') if assignment_folder else config.OUTPUT_DIR

    # Setup session with CSRF
    cookies_list = json.load(open(config.COOKIES_FILE))
    s = req.Session()
    csrf = None
    for c in cookies_list:
        s.cookies.set(c['name'], c['value'], domain=c['domain'], path=c['path'])
        if c['name'] == '_csrf_token':
            csrf = req.utils.unquote(c['value'])

    def hdrs(ct='application/json', use_csrf=True):
        h = {'Accept': 'application/json+canvas-string-ids', 'X-Requested-With': 'XMLHttpRequest', 'User-Agent': 'Mozilla/5.0'}
        if ct: h['Content-Type'] = ct
        if use_csrf and csrf: h['X-CSRF-Token'] = csrf
        return h

    # Find files
    files = [f for f in Path(output_dir).iterdir()
             if f.is_file() and not f.name.startswith('.') and f.suffix in ['.docx','.pdf','.png','.jpg']]
    if not files:
        log.error("No files found")
        return False

    log.info(f"Uploading {len(files)} file(s)...")
    file_ids = []

    for p in files:
        log.info(f"  {p.name} ({os.path.getsize(p)} bytes)")
        try:
            # 1. Request token
            r = s.post('https://psu.instructure.com/files/pending', data={
                'name': p.name, 'on_duplicate': 'rename', 'no_redirect': 'true',
                'attachment[filename]': p.name, 'attachment[size]': str(os.path.getsize(p)),
                'attachment[context_code]': f'course_{course_id}',
                'attachment[content_type]': mimetypes.guess_type(str(p))[0] or 'application/octet-stream'
            }, headers=hdrs('application/x-www-form-urlencoded'))
            r.raise_for_status()
            info = r.json()

            # 2. Upload
            upload_url = info['upload_url']
            if info['upload_params'].get('token'):
                upload_url += f"?token={info['upload_params']['token']}"
            with open(p, 'rb') as f:
                r = s.post(upload_url, files={'file': (p.name, f)}, headers=hdrs(None, False), allow_redirects=False)
            r.raise_for_status()

            # 3. Confirm
            r = s.get(r.headers['location'], headers=hdrs())
            r.raise_for_status()
            fid = str(r.json()['id'])
            file_ids.append(fid)
            log.info(f"    ID: {fid}")
            time.sleep(0.5)

        except Exception as e:
            log.error(f"    Failed: {e}")

    if not file_ids:
        log.error("No files uploaded")
        return False

    # 4. Submit
    log.info(f"Submitting {len(file_ids)} file(s)...")
    try:
        r = s.post(f"https://psu.instructure.com/courses/{course_id}/assignments/{assign_id}/submissions",
            data={'submission[submission_type]': 'online_upload', 'submission[attachment_ids]': ','.join(file_ids), 'authenticity_token': csrf},
            headers=hdrs('application/x-www-form-urlencoded'))
        r.raise_for_status()
        log.info(f"Success! ID: {r.json().get('id')}")
        return True
    except Exception as e:
        log.error(f"Submit failed: {e}")
        return False


def run_gui(url, product, model, prompt, ref_files=[], assignment_folder=None, progress=None):
    """GUI entry point"""
    if not assignment_folder:
        raise ValueError("assignment_folder required")

    output_dir = os.path.join(assignment_folder, 'auto', 'output')
    os.makedirs(output_dir, exist_ok=True)
    for f in Path(output_dir).glob('*'):
        if f.is_file(): f.unlink()

    if progress: progress.update(status="Fetching...", progress=10)
    data = get_homework_details(url)
    desc = unescape(re.sub(r'<[^>]+>', '\n', data['assignment']['description'])).strip()

    if progress: progress.update(status="Asking LLM...", progress=30)
    answer = ask_llm_with_pdfs(desc, product, model, prompt, ref_files)
    if '[no]' in answer.lower()[:100]:
        raise Exception("LLM returned [no]")

    if progress: progress.update(status="Generating output...", progress=70)
    clean, imgs = parse_img_requests(answer)
    answer_path = os.path.join(output_dir, 'answer.md')
    with open(answer_path, 'w', encoding='utf-8') as f:
        f.write(clean)
    generate_images(imgs, output_dir)

    docx_path = os.path.join(output_dir, 'answer.docx')
    md_to_docx(docx_path, answer_path)

    if progress: progress.update(status="Done", progress=100)
    return {'status': 'success', 'answer_path': answer_path, 'docx_path': docx_path}


def main(url=None, product=None, model=None):
    """CLI entry point"""
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('--url', type=str)
    parser.add_argument('--product', choices=['Gemini','Claude'])
    parser.add_argument('--model', type=str)
    args = parser.parse_args()

    url = args.url or url or TARGET_URL
    product = args.product or product or 'Gemini'
    model = args.model or model or 'gemini-2.5-pro'

    log.info(f"URL: {url}")
    log.info(f"Model: {product} {model}")

    config.ensure_dirs()
    for f in Path(config.OUTPUT_DIR).glob('*'):
        if f.is_file(): f.unlink()

    data = get_homework_details(url)
    desc = unescape(re.sub(r'<[^>]+>', '\n', data['assignment']['description'])).strip()

    answer = ask_llm_with_pdfs(desc, product, model, config.DEFAULT_PROMPTS['homework'])
    if '[no]' in answer.lower()[:100]:
        log.error("LLM returned [no]")
        return

    clean, imgs = parse_img_requests(answer)
    answer_path = os.path.join(config.OUTPUT_DIR, 'answer.md')
    with open(answer_path, 'w', encoding='utf-8') as f:
        f.write(clean)
    generate_images(imgs, config.OUTPUT_DIR)
    md_to_docx(os.path.join(config.OUTPUT_DIR, 'answer.docx'), answer_path)

    log.info(f"Done! Files in: {config.OUTPUT_DIR}")

    if input("Submit? (y/n): ").lower() == 'y':
        submit_to_canvas(url)


if __name__ == '__main__':
    main()
