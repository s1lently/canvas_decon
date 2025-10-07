import requests, json, os, sys, re, mimetypes, time
from urllib.parse import urlparse
from html import unescape
import google.generativeai as genai
from pathlib import Path

# Configuration
TARGET_ASSIGNMENT_URL = "https://psu.instructure.com/courses/2418560/assignments/17474475"
COOKIES_FILE, OUTPUT_DIR, SUBMISSION_DIR = 'cookies.json', 'homework_res', 'homework_res/submission'
DETAILS_FILE, ANSWER_FILE = f'{OUTPUT_DIR}/assignment_details.json', f'{OUTPUT_DIR}/answer.md'
GEMINI_API_KEY, PDF_DIR = "AIzaSyBZTx5UDH7pxyYZUgpDzKHRU25FWoPIA8I", 'bisc_pdfs'
PERSONAL_INFO_FILE = 'personal_info.json'

def get_homework_details():
    path = urlparse(TARGET_ASSIGNMENT_URL).path.split('/')
    course_id, assignment_id = path[path.index('courses')+1], path[path.index('assignments')+1]
    
    cookies = {c['name']: c['value'] for c in json.load(open(COOKIES_FILE))}
    session = requests.Session()
    session.cookies.update(cookies)
    session.headers.update({'User-Agent': 'Mozilla/5.0', 'Accept': 'application/json+canvas-string-ids'})
    
    assignment = session.get(f"https://psu.instructure.com/api/v1/courses/{course_id}/assignments/{assignment_id}", timeout=20).json()
    submission = session.get(f"https://psu.instructure.com/api/v1/courses/{course_id}/assignments/{assignment_id}/submissions/self", timeout=20).json()
    
    data = {'assignment': {k: assignment.get(k) for k in ['id', 'name', 'description', 'due_at', 'points_possible', 'submission_types', 'html_url']},
            'submission': {k: submission.get(k) for k in ['id', 'submitted_at', 'workflow_state', 'score', 'grade', 'late', 'missing']}}
    json.dump(data, open(DETAILS_FILE, 'w', encoding='utf-8'), indent=2, ensure_ascii=False)
    print(f"✓ Assignment: {assignment['name']}")
    return data

def ask_llm_with_pdfs(description):
    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel('gemini-2.5-pro')
    
    pdf_files = [genai.upload_file(path=str(p), display_name=p.name) or print(f"📄 Uploading: {p.name}") for p in Path(PDF_DIR).glob('*.pdf')] if Path(PDF_DIR).exists() else []
    
    # Load personal info
    try:
        personal_info = json.load(open(PERSONAL_INFO_FILE))
        personal_context = f"""
**Your Personal Information (USE THESE EXACT VALUES):**
- Name: {personal_info['name']}
- Age: {personal_info['age']} years
- Weight: {personal_info['weight_kg']} kg ({personal_info['weight_lbs']} lbs)
- Height: {personal_info['height_cm']} cm ({personal_info['height_inches']} inches)
- Gender: {personal_info['gender']}
- Location: {personal_info['location']}
"""
    except:
        personal_context = ""
    
    prompt = f"""Analyze the description and complete the assignment:
{personal_context}

**Decision Logic:**
1. If requires course materials (lecture/textbook) → check PDF attachments for relevant content → [yes] continue
2. If requires personal experience/real-life examples (food labels, daily life) → no attachments needed → directly [yes] and write content

**Answer Requirements:**
1. Pure English
2. Avoid complex sentences/vocabulary. Avoid phrases like "resulting in", "including", "leading to". Keep logic simple. Only use necessary technical terms.
3. Start with status: [yes] = can complete [no] = cannot complete (only if description incomplete)
4. Use clean markdown format

**Format Rules (strictly follow):**
- ❌ No dividers: ----, ====, ———
- ❌ No code blocks: ```
- ❌ No decorative symbols: tables, frames, ASCII art
- ❌ No bullet points: *, -, •
- ✅ Bold question numbers: **1)** **2)** **3)** or **Question 1:**
- ✅ Letter numbering for multiple points: a. \\n b. \\n c. ... (each on new line)
- ✅ Only use: headers (###), bold (**), paragraphs, letter numbering

**Image Request Annotation:**
ONLY if assignment EXPLICITLY requires uploading an image/photo/file (e.g., "Upload a copy of the food label", "Attach a photo"), add at the **end** of answer:
[gen_img]
{{
name: filename.png
des: Detailed description for image generation. IMPORTANT: Infer the full context from the question and describe what the image should realistically show based on the assignment requirements.
}}
DO NOT generate images for text-only calculation questions or written reflections.

**Personal Experience Rules:**
- Location: Middletown, PA
- Background: AI CV, C++, Linux, reverse engineering (only reference if needed)
- Food/daily life: Write common realistic scenarios naturally

**Description:**
{description}

**Note:** For personal experience assignments, write naturally like a real student. Make question numbers prominent.
"""
    
    print("\n🤖 Asking LLM...")
    return model.generate_content([prompt] + pdf_files).text

def parse_img_requests(answer_text):
    img_requests = []
    for match in re.finditer(r'\[gen_img\]\s*\{([^}]+)\}', answer_text, re.DOTALL):
        content = match.group(1)
        name_match = re.search(r'name:\s*(.+?)(?:\n|$)', content)
        des_match = re.search(r'des:\s*(.+?)(?:\n|$)', content, re.DOTALL)
        if name_match and des_match:
            img_requests.append({'name': name_match.group(1).strip(), 'description': des_match.group(1).strip()})
    return re.sub(r'\[gen_img\]\s*\{[^}]+\}', '', answer_text, flags=re.DOTALL).strip(), img_requests

def generate_images(img_requests):
    if not img_requests: return
    print(f"\n🖼️  Generating {len(img_requests)} image(s) with Gemini...")
    from google import genai as genai_client
    from PIL import Image
    from io import BytesIO
    
    client = genai_client.Client(api_key=GEMINI_API_KEY)
    for req in img_requests:
        print(f"  - {req['name']}: {req['description'][:50]}...")
        try:
            response = client.models.generate_content(model="gemini-2.5-flash-image", contents=[req['description']])
            for part in response.candidates[0].content.parts:
                if part.inline_data:
                    Image.open(BytesIO(part.inline_data.data)).save(img_path := f"{SUBMISSION_DIR}/{req['name']}")
                    print(f"    ✅ Generated: {img_path}")
                    break
        except Exception as e: print(f"    ❌ Failed: {e}")

def md_to_docx(output_path):
    from docx import Document
    from docx.shared import Inches
    doc = Document()
    for line in open(ANSWER_FILE, 'r', encoding='utf-8'):
        line = line.rstrip()
        if line.strip() in ['[yes]', '[no]'] or line.startswith('```'): continue
        if line.startswith('### '): doc.add_heading(line[4:], 3)
        elif line.startswith('## '): doc.add_heading(line[3:], 2)
        elif line.startswith('# '): doc.add_heading(line[2:], 1)
        elif line.startswith(('*   ', '- ')):
            text = line[4:] if line.startswith('*   ') else line[2:]
            p = doc.add_paragraph(style='List Bullet')
            if '**' in text:
                for i, part in enumerate(text.split('**')): (run := p.add_run(part)).bold = bool(i % 2)
            else: p.add_run(text)
        elif line.startswith('    '):
            p = doc.add_paragraph()
            for i, part in enumerate(line.strip().split('**')): (run := p.add_run(part)).bold = bool(i % 2)
            p.paragraph_format.left_indent = Inches(0.5)
        elif line.strip():
            p = doc.add_paragraph()
            for i, part in enumerate(line.split('**')): (run := p.add_run(part)).bold = bool(i % 2)
    doc.save(output_path)
    print(f"✅ Converted to: {output_path}")

def submit_to_canvas():
    print("\n📤 Step 7: Submitting to Canvas...")
    path = urlparse(TARGET_ASSIGNMENT_URL).path.split('/')
    course_id, assignment_id = path[path.index('courses')+1], path[path.index('assignments')+1]
    
    cookies_list = json.load(open(COOKIES_FILE))
    session = requests.Session()
    csrf_token = None
    for c in cookies_list:
        session.cookies.set(c['name'], c['value'], domain=c['domain'], path=c['path'])
        if c['name'] == '_csrf_token': csrf_token = requests.utils.unquote(c['value'])
    
    def update_csrf():
        nonlocal csrf_token
        if '_csrf_token' in session.cookies and (new := requests.utils.unquote(session.cookies['_csrf_token'])): csrf_token = new
    
    def headers(ct='application/json', csrf=True):
        h = {'Accept': 'application/json+canvas-string-ids, application/json, text/plain, */*',
             'Referer': f'https://psu.instructure.com/courses/{course_id}/assignments/{assignment_id}',
             'X-Requested-With': 'XMLHttpRequest', 'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36'}
        if ct: h['Content-Type'] = ct
        if csrf: h['X-CSRF-Token'] = csrf_token
        return h
    
    files = [f for f in Path(SUBMISSION_DIR).iterdir() if f.is_file() and not f.name.startswith('.')]
    if not files: return print("❌ No files found") or False
    
    print(f"Files: {[f.name for f in files]}")
    file_ids = []
    for p in files:
        print(f"\n  {'='*50}\n  {p.name} ({os.path.getsize(p)} bytes)")
        try:
            print(f"  [1/4] Request token...")
            r = session.post('https://psu.instructure.com/files/pending', 
                           data={'name': p.name, 'on_duplicate': 'rename', 'no_redirect': 'true',
                                 'attachment[intent]': 'submit', 'attachment[asset_string]': f'assignment_{assignment_id}',
                                 'attachment[filename]': p.name, 'attachment[size]': str(os.path.getsize(p)),
                                 'attachment[context_code]': f'course_{course_id}', 'attachment[on_duplicate]': 'rename',
                                 'attachment[content_type]': mimetypes.guess_type(str(p))[0] or 'application/octet-stream'},
                           headers=headers('application/x-www-form-urlencoded'))
            r.raise_for_status()
            update_csrf()
            info = r.json()
            print(f"  ✓ Got token") or time.sleep(0.5)
            
            print(f"  [2/4] Upload...")
            url = f"{info['upload_url']}?token={info['upload_params'].get('token')}" if info['upload_params'].get('token') else info['upload_url']
            f = open(p, 'rb')
            r = session.post(url, files={'file': (p.name, f, info['upload_params'].get('content_type'))}, headers=headers(None, False), allow_redirects=False)
            f.close()
            r.raise_for_status()
            print(f"  ✓ Uploaded") or time.sleep(0.5)
            
            print(f"  [3/4] Confirm...")
            r = session.get(r.headers.get('location'), headers=headers())
            r.raise_for_status()
            fid = str(r.json()['id'])
            update_csrf()
            file_ids.append(fid)
            print(f"  ✓ File ID: {fid}") or time.sleep(1)
        except Exception as e: print(f"  ❌ Failed: {e}")
    
    if not file_ids: return print("\n❌ No files uploaded") or False
    
    print(f"\n  [4/4] Submit {len(file_ids)} file(s)...")
    data = {'utf8': '✓', 'authenticity_token': csrf_token, 'submission[submission_type]': 'online_upload',
            'submission[attachment_ids]': ','.join(file_ids), 'submission[eula_agreement_timestamp]': '', 'submission[comment]': ''}
    for i, fid in enumerate(file_ids): data[f'attachments[{i}][uploaded_data]'] = f'C:\\fakepath\\f{i}.dat'
    
    try:
        r = session.post(f"https://psu.instructure.com/courses/{course_id}/assignments/{assignment_id}/submissions", data=data, headers=headers('application/x-www-form-urlencoded; charset=UTF-8'))
        r.raise_for_status()
        print(f"  ✅ Success! ID: {r.json().get('id')}")
        return True
    except Exception as e: print(f"  ❌ Failed: {e}") or False

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(SUBMISSION_DIR, exist_ok=True)
    
    # Clean submission folder from previous runs
    for f in Path(SUBMISSION_DIR).glob('*'):
        if f.is_file(): f.unlink()
    
    print("Step 1: Fetching assignment details...")
    data = get_homework_details()
    
    print("\nStep 2: Extracting description...")
    description = unescape(re.sub(r'<[^>]+>', '\n', data['assignment']['description'])).strip()
    print(f"Description:\n{description}\n")
    
    print("Step 3: Asking LLM with PDFs...")
    answer = ask_llm_with_pdfs(description)
    
    print("\nStep 4: Validating and parsing answer...")
    if '[no]' in answer.lower()[:100]: sys.exit(print("❌ LLM returned [no] - Exiting..."))
    
    clean_answer, img_requests = parse_img_requests(answer)
    open(ANSWER_FILE, 'w', encoding='utf-8').write(clean_answer)
    print(f"✓ Answer saved to: {ANSWER_FILE}")
    
    print("\nStep 5: Generating requested images...")
    generate_images(img_requests)
    
    print("\nStep 6: Converting to DOCX...")
    md_to_docx(f"{SUBMISSION_DIR}/answer.docx")
    
    print(f"\n✅ All steps completed! Files in: {SUBMISSION_DIR}/")
    
    # Submit to Canvas
    print("\n" + "="*60)
    user_input = input("📤 Submit to Canvas? (yes/no): ").strip().lower()
    if user_input in ['yes', 'y']:
        submit_to_canvas()
        print("\n" + "="*60)
        print("🎉 All done!")
    else:
        print("⏸️  Submission skipped.")

if __name__ == '__main__':
    main()
